import json
import sys
from pathlib import Path


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                blocks.append(text)
    return "\n".join(blocks)


def read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()


def main() -> None:
    path = Path(sys.argv[1])
    suffix = path.suffix.lower()
    if suffix == ".docx":
        text = read_docx(path)
    elif suffix == ".pdf":
        text = read_pdf(path)
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")
    print(json.dumps({"text": text}, ensure_ascii=False))


if __name__ == "__main__":
    main()
